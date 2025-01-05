from flask import Blueprint, request, jsonify, send_file, current_app
from src.services import DocumentService, TranslationService
from src.models import Book, Fragment
from src.database import db
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from config.config import Config

book_api = Blueprint('book_api', __name__)
document_service = DocumentService()
translation_service = TranslationService()

# 添加翻译进度存储
translation_progress = {}

@book_api.route('/books', methods=['GET'])
def get_books():
    """获取所有书籍"""
    books = Book.query.all()
    return jsonify([book.to_dict() for book in books])

@book_api.route('/books/<int:book_id>', methods=['GET'])
def get_book(book_id):
    """获取单本书籍信息"""
    try:
        book = Book.query.get_or_404(book_id)
        return jsonify(book.to_dict())
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@book_api.route('/upload', methods=['POST'])
def upload_file():
    """上传文件"""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件上传'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
    
    if not document_service.allowed_file(file.filename):
        return jsonify({'error': '不支持的文件类型'}), 400

    try:
        # 保存文件
        filepath = document_service.save_file(file)
        
        # 处理文档
        text = document_service.extract_text(filepath)
        fragments = document_service.split_text(text)
        
        # 保存到数据库
        book = Book(name=file.filename, total_fragments=len(fragments))
        db.session.add(book)
        db.session.commit()
        
        for idx, content in enumerate(fragments):
            fragment = Fragment(
                book_id=book.id,
                fragment_number=idx + 1,
                original_text=content,
                translated_text=""
            )
            db.session.add(fragment)
        
        db.session.commit()
        return jsonify({'success': True, 'book_id': book.id})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@book_api.route('/book/<int:book_id>/fragments', methods=['GET'])
def get_fragments(book_id):
    """获取书籍的所有文本碎片"""
    fragments = Fragment.query.filter_by(book_id=book_id).all()
    return jsonify([fragment.to_dict() for fragment in fragments])

@book_api.route('/translation_progress/<int:book_id>', methods=['GET'])
def get_translation_progress(book_id):
    """获取翻译进度"""
    progress = translation_progress.get(book_id, {'completed': 0, 'total': 0})
    return jsonify(progress)

@book_api.route('/translation_settings/<int:book_id>', methods=['GET'])
def get_translation_settings(book_id):
    """获取书籍的翻译设置"""
    try:
        book = Book.query.get_or_404(book_id)
        
        # 检查是否有已翻译的碎片
        translated_count = Fragment.query.filter(
            Fragment.book_id == book_id,
            Fragment.translated_text != None,
            Fragment.translated_text != ""
        ).count()
        
        return jsonify({
            'has_settings': bool(book.target_language and book.translation_style and book.prompt_template),
            'settings': {
                'target_language': book.target_language,
                'translation_style': book.translation_style,
                'prompt_template': book.prompt_template,
                'custom_prompt': book.custom_prompt,
                'has_translated_fragments': translated_count > 0
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@book_api.route('/save_translation_settings/<int:book_id>', methods=['POST'])
def save_translation_settings(book_id):
    """保存书籍的翻译设置"""
    try:
        book = Book.query.get_or_404(book_id)
        
        # 检查是否有已翻译的碎片
        translated_count = Fragment.query.filter(
            Fragment.book_id == book_id,
            Fragment.translated_text != None,
            Fragment.translated_text != ""
        ).count()
        
        if translated_count > 0:
            return jsonify({
                'success': False,
                'error': '已有翻译完成的碎片，无法修改翻译设置'
            }), 400
        
        # 更新翻译设置
        settings = request.json
        book.target_language = settings.get('target_language', '中文')
        book.translation_style = settings.get('translation_style', '准确、流畅')
        book.prompt_template = settings.get('prompt_template', 'standard')
        book.custom_prompt = settings.get('custom_prompt', '')
        
        db.session.commit()
        return jsonify({
            'success': True,
            'message': '翻译设置已保存'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@book_api.route('/translate_remaining/<int:book_id>', methods=['POST'])
def translate_remaining(book_id):
    """翻译书籍中所有未翻译的碎片"""
    try:
        book = Book.query.get_or_404(book_id)
        
        # 检查是否有翻译设置
        if not (book.target_language and book.translation_style and book.prompt_template):
            return jsonify({'error': '请先设置翻译参数'}), 400
            
        # 修改查询条件，同时匹配 None 和空字符串
        fragments = Fragment.query.filter(
            Fragment.book_id == book_id,
            (Fragment.translated_text == None) | (Fragment.translated_text == "")  # noqa
        ).all()
        
        total = len(fragments)
        if total == 0:
            return jsonify({'message': '没有需要翻译的内容'}), 200
            
        translation_progress[book_id] = {'completed': 0, 'total': total, 'errors': []}
        
        for idx, fragment in enumerate(fragments, 1):
            try:
                # 使用书籍的翻译设置进行翻译
                translated_text = translation_service.translate(
                    fragment.original_text,
                    target_lang=book.target_language,
                    style=book.translation_style,
                    custom_prompt=book.custom_prompt if book.custom_prompt else None
                )
                
                if translated_text:  # 只有在成功获取译文时才更新
                    fragment.translated_text = translated_text
                    db.session.add(fragment)
                    # 立即提交每个翻译结果
                    db.session.commit()
                    
                    # 更新进度信息
                    translation_progress[book_id]['completed'] = idx
                    
                    # 更新书籍的翻译进度
                    total_fragments = Fragment.query.filter_by(book_id=book_id).count()
                    translated_fragments = Fragment.query.filter(
                        Fragment.book_id == book_id,
                        Fragment.translated_text != None,
                        Fragment.translated_text != ""
                    ).count()
                    book.progress = round((translated_fragments / total_fragments) * 100)
                    db.session.add(book)
                    db.session.commit()
                
            except Exception as e:
                error_msg = f"翻译碎片 #{fragment.fragment_number} 失败: {str(e)}"
                current_app.logger.error(error_msg)
                translation_progress[book_id]['errors'].append(error_msg)
                continue
        
        # 检查是否有错误发生
        if translation_progress[book_id].get('errors'):
            error_summary = "\n".join(translation_progress[book_id]['errors'])
            return jsonify({
                'warning': '部分内容翻译失败',
                'details': error_summary,
                'completed': translation_progress[book_id]['completed'],
                'total': total
            }), 206  # Partial Content
        
        return jsonify({'success': True})
        
    except Exception as e:
        db.session.rollback()
        error_msg = f"翻译过程出错: {str(e)}"
        current_app.logger.error(error_msg)
        return jsonify({'error': error_msg}), 500
    finally:
        # 清理进度信息
        translation_progress.pop(book_id, None)

@book_api.route('/retranslate/<int:book_id>', methods=['POST'])
def retranslate_book(book_id):
    """重新翻译整本书"""
    try:
        book = Book.query.get_or_404(book_id)
        
        # 删除所有现有碎片
        Fragment.query.filter_by(book_id=book_id).delete()
        db.session.commit()
        
        # 重新处理文档
        filepath = os.path.join(Config.UPLOAD_FOLDER, book.name)
        if not os.path.exists(filepath):
            return jsonify({'error': '原文件不存在'}), 404
            
        text = document_service.extract_text(filepath)
        fragments = document_service.split_text(text)
        
        # 更新总碎片数
        book.total_fragments = len(fragments)
        translation_progress[book_id] = {'completed': 0, 'total': len(fragments)}
        
        # 创建新碎片并翻译
        for idx, content in enumerate(fragments):
            try:
                fragment = Fragment(
                    book_id=book.id,
                    fragment_number=idx + 1,
                    original_text=content,
                    translated_text=translation_service.translate(content)
                )
                db.session.add(fragment)
                # 立即提交每个翻译结果
                db.session.commit()
                
                # 更新进度
                translation_progress[book_id]['completed'] = idx + 1
                
                # 更新书籍的翻译进度
                book.progress = round(((idx + 1) / len(fragments)) * 100)
                db.session.add(book)
                db.session.commit()
                
            except Exception as e:
                # 如果单个片段翻译失败，记录错误但继续处理其他片段
                current_app.logger.error(f"重新翻译片段 {idx + 1} 时出错: {str(e)}")
                continue
        
        # 清理进度信息
        translation_progress.pop(book_id, None)
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        # 清理进度信息
        translation_progress.pop(book_id, None)
        return jsonify({'error': str(e)}), 500

@book_api.route('/export/<int:book_id>', methods=['GET'])
def export_book(book_id):
    """导出书籍的原文和译文"""
    try:
        # 获取书籍信息
        book = Book.query.get_or_404(book_id)
        fragments = Fragment.query.filter_by(book_id=book_id).order_by(Fragment.fragment_number).all()
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(book.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加内容
        for fragment in fragments:
            if fragment.translated_text:  # 只导出已翻译的部分
                # 添加原文
                p = doc.add_paragraph()
                p.add_run('原文：').bold = True
                p.add_run(fragment.original_text)
                
                # 添加译文
                p = doc.add_paragraph()
                p.add_run('译文：').bold = True
                p.add_run(fragment.translated_text)
                
                # 添加分隔线
                doc.add_paragraph('=' * 50)
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # 发送文件
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f"{book.name}_翻译.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@book_api.route('/delete/<int:book_id>', methods=['DELETE'])
def delete_book(book_id):
    """删除书籍及其所有碎片"""
    try:
        book = Book.query.get_or_404(book_id)
        
        # 删除上传的原文件
        filepath = os.path.join(Config.UPLOAD_FOLDER, book.name)
        if os.path.exists(filepath):
            os.remove(filepath)
        
        # 删除书籍（级联删除会自动删除相关的碎片）
        db.session.delete(book)
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500 

@book_api.route('/untranslated_fragments/<int:book_id>', methods=['GET'])
def get_untranslated_fragments(book_id):
    """获取未翻译的片段"""
    try:
        book = Book.query.get_or_404(book_id)
        fragments = Fragment.query.filter(
            Fragment.book_id == book_id,
            (Fragment.translated_text == None) | (Fragment.translated_text == "")  # noqa
        ).all()
        return jsonify([{
            'id': f.id,
            'content': f.original_text
        } for f in fragments])
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@book_api.route('/translate_fragment/<int:book_id>/<int:fragment_id>', methods=['POST'])
def translate_fragment(book_id, fragment_id):
    """翻译单个片段"""
    try:
        # 获取书籍和片段
        book = Book.query.get_or_404(book_id)
        fragment = Fragment.query.filter_by(id=fragment_id, book_id=book_id).first_or_404()
        
        # 如果已经翻译过，直接返回成功
        if fragment.translated_text is not None and fragment.translated_text != "":
            return jsonify({'message': '片段已翻译'}), 200
        
        # 获取翻译设置
        settings_response = get_translation_settings(book_id)
        if settings_response.status_code != 200:
            return settings_response
        
        settings = settings_response.get_json()
        if not settings.get('has_settings'):
            return jsonify({'error': '请先设置翻译参数'}), 400
        
        # 执行翻译
        translated_text = translation_service.translate(
            fragment.original_text,
            target_lang=settings['settings']['target_language'],
            style=settings['settings']['translation_style'],
            custom_prompt=settings['settings'].get('custom_prompt')
        )
        
        # 保存翻译结果
        fragment.translated_text = translated_text
        db.session.commit()
        
        return jsonify({'message': '翻译成功'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400 